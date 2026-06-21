from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "SAD.md"
OUTPUT = ROOT / "docs" / "SAD_Arquitectura_Aplicaciones_Grupo11.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.text = "SAD - Arquitectura de Aplicaciones - Grupo 11"
    header.style = styles["Header"]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    footer = section.footer.paragraphs[0]
    footer.text = "Trabajo Practico - Primer cuatrimestre 2026"
    footer.style = styles["Footer"]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_title_page(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Software Architecture Document (SAD)")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Ecosistema de microservicios UADE - Grupo 11")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string("1F4D78")

    meta = doc.add_table(rows=5, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.autofit = False
    rows = [
        ("Materia", "Arquitectura de Aplicaciones"),
        ("Cuatrimestre", "Primer cuatrimestre de 2026"),
        ("Version", "1.0"),
        ("Fecha", "8 de junio de 2026"),
        ("Alcance", "PoC con RabbitMQ, Eureka, Gateway, JWT y order-service"),
    ]
    for row, (label, value) in zip(meta.rows, rows):
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_width(row.cells[0], 2300)
        set_cell_width(row.cells[1], 6500)
        for cell in row.cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(row.cells[0], "E8EEF5")
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = note.add_run("Documento generado a partir del repositorio y de la PoC validada con Docker Compose.")
    r.italic = True
    r.font.color.rgb = RGBColor.from_string("555555")
    doc.add_page_break()


def add_toc(doc, headings):
    doc.add_heading("Indice", level=1)
    for level, text in headings:
        if level <= 2:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25 * (level - 1))
            p.add_run(text)
    doc.add_page_break()


def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        r.font.size = Pt(9)


def add_table(doc, table_lines):
    rows = []
    for line in table_lines:
        if not line.strip().startswith("|"):
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", c) for c in cells):
            continue
        rows.append(cells)

    if not rows:
        return
    col_count = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [int(9360 / col_count)] * col_count
    for i, row in enumerate(rows):
        for j, text in enumerate(row):
            cell = table.cell(i, j)
            cell.text = text
            set_cell_width(cell, widths[j])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if i == 0:
                set_cell_shading(cell, "F2F4F7")
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    doc.add_paragraph()


def add_paragraph_with_inline_code(doc, text):
    p = doc.add_paragraph()
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1])
            r.font.name = "Consolas"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor.from_string("1F4D78")
        else:
            p.add_run(part)


def build():
    markdown = SOURCE.read_text(encoding="utf-8").splitlines()
    headings = []
    for line in markdown:
        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            text = line.lstrip("#").strip()
            headings.append((level, text))

    doc = Document()
    style_document(doc)
    add_title_page(doc)
    add_toc(doc, headings)

    in_code = False
    code_lines = []
    table_lines = []

    def flush_table():
        nonlocal table_lines
        if table_lines:
            add_table(doc, table_lines)
            table_lines = []

    for raw in markdown:
        line = raw.rstrip()

        if line == "---":
            continue
        if line.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                flush_table()
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue

        if line.startswith("|"):
            table_lines.append(line)
            continue
        flush_table()

        if not line.strip():
            continue
        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            text = line.lstrip("#").strip()
            if level == 1 and "Software Architecture Document" in text:
                continue
            doc.add_heading(text, level=min(level, 3))
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(line[2:])
            continue
        if re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            p.add_run(re.sub(r"^\d+\. ", "", line))
            continue
        add_paragraph_with_inline_code(doc, line)

    flush_table()
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
